"""
Generate SUBSCRIPTION_DATA_FLOW.docx from the markdown file.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import copy

# ─── Colours ──────────────────────────────────────────────────────────────────
BLUE_DARK   = RGBColor(0x1F, 0x38, 0x64)   # heading 1
BLUE_MID    = RGBColor(0x2E, 0x54, 0x96)   # heading 2/3
BLUE_LIGHT  = RGBColor(0x21, 0x63, 0xAB)   # heading 4
BLACK       = RGBColor(0x00, 0x00, 0x00)
GRAY_BG     = "F2F2F2"   # code block background
HEADER_BG   = "D5E8F0"   # table header background
ALT_ROW     = "FAFAFA"   # alternate row (subtle)
NOTE_BG     = "FFF8DC"   # blockquote / note background

# ─── Helper: set cell background ─────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

# ─── Helper: set paragraph border (bottom only) ──────────────────────────────
def add_bottom_border(para, color="2E5496", size=6):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(size))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

# ─── Helper: parse inline markdown → list of (text, bold, italic, code) ──────
def parse_inline(text):
    """Return list of (fragment_text, bold, italic, is_code)."""
    segments = []
    # Combined regex for **bold**, *italic*, `code`
    pattern = re.compile(r'(`[^`]+`|\*\*[^*]+\*\*|\*[^*]+\*)')
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            segments.append((text[pos:m.start()], False, False, False))
        tok = m.group(0)
        if tok.startswith('`'):
            segments.append((tok[1:-1], False, False, True))
        elif tok.startswith('**'):
            segments.append((tok[2:-2], True, False, False))
        else:
            segments.append((tok[1:-1], False, True, False))
        pos = m.end()
    if pos < len(text):
        segments.append((text[pos:], False, False, False))
    return segments

def add_inline(para, text, base_size=10.5, base_color=BLACK, base_bold=False):
    """Add inline-formatted text to an existing paragraph."""
    for frag, bold, italic, is_code in parse_inline(text):
        if not frag:
            continue
        run = para.add_run(frag)
        run.bold = bold or base_bold
        run.italic = italic
        if is_code:
            run.font.name = 'Courier New'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        else:
            run.font.name = 'Arial'
            run.font.size = Pt(base_size)
            run.font.color.rgb = base_color

# ─── Helper: add a code block (monospace, gray bg) ───────────────────────────
def add_code_block(doc, lines):
    """Add a shaded monospace block for ASCII / code content."""
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.left_indent  = Cm(0.5)
        # shade paragraph background
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), GRAY_BG)
        pPr.append(shd)
        run = p.add_run(line if line else ' ')
        run.font.name = 'Courier New'
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    # small gap after block
    gap = doc.add_paragraph()
    gap.paragraph_format.space_before = Pt(2)
    gap.paragraph_format.space_after  = Pt(2)

# ─── Helper: add note / blockquote ───────────────────────────────────────────
def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.8)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), NOTE_BG)
    pPr.append(shd)
    add_inline(p, text.lstrip('> '), base_size=10)

# ─── Helper: add a markdown table ────────────────────────────────────────────
def add_md_table(doc, headers, rows, col_widths_cm):
    n = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # set column widths
    for i, w in enumerate(col_widths_cm):
        for row in table.rows:
            row.cells[i].width = Cm(w)

    # header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, HEADER_BG)
        p = cell.paragraphs[0]
        add_inline(p, h, base_size=10, base_bold=True)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # data rows
    for ri, row_data in enumerate(rows):
        tr = table.rows[ri + 1]
        bg = "FFFFFF" if ri % 2 == 0 else ALT_ROW
        for ci, cell_text in enumerate(row_data):
            cell = tr.cells[ci]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            add_inline(p, cell_text, base_size=9.5)

    doc.add_paragraph()  # spacing after table

# ─── Helper: heading helpers ─────────────────────────────────────────────────
def h1(doc, text):
    p = doc.add_heading(level=1)
    p.clear()
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = BLUE_DARK
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(8)
    add_bottom_border(p, "1F3864", 12)
    return p

def h2(doc, text):
    p = doc.add_heading(level=2)
    p.clear()
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = BLUE_MID
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    add_bottom_border(p, "2E5496", 6)
    return p

def h3(doc, text):
    p = doc.add_heading(level=3)
    p.clear()
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = BLUE_MID
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    return p

def h4(doc, text):
    p = doc.add_heading(level=4)
    p.clear()
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = BLUE_LIGHT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    return p

def body(doc, text, bold=False, size=10.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    add_inline(p, text, base_size=size, base_bold=bold)
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    add_inline(p, text, base_size=10)
    return p

def divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)

# ─── Build document ───────────────────────────────────────────────────────────
doc = Document()

# Page setup: A4, narrow margins
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2)
section.right_margin  = Cm(2)
section.top_margin    = Cm(2)
section.bottom_margin = Cm(2)

# Default paragraph font
doc.styles['Normal'].font.name = 'Arial'
doc.styles['Normal'].font.size = Pt(10.5)

# ══════════════════════════════════════════════════════════════════════════════
# COVER / TITLE
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(30)
p.paragraph_format.space_after  = Pt(6)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Luồng Dữ Liệu Subscription')
run.font.name = 'Arial'
run.font.size = Pt(24)
run.font.bold = True
run.font.color.rgb = BLUE_DARK

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(4)
run2 = p2.add_run('Khách Hàng Phổ Thông & Đại Lý')
run2.font.name = 'Arial'
run2.font.size = Pt(18)
run2.font.color.rgb = BLUE_MID

doc.add_paragraph()

meta_lines = [
    ('Dự án:', 'billing-CA'),
    ('Cập nhật:', '2026-05-20'),
    ('Phạm vi:', 'Mô tả toàn bộ các bảng liên quan đến hai loại khách hàng, mối quan hệ giữa các bảng, luồng nghiệp vụ tạo subscription, và dữ liệu đi đâu sau khi subscription được tạo.'),
]
for label, val in meta_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(label + '  ')
    r1.font.bold = True
    r1.font.name = 'Arial'
    r1.font.size = Pt(11)
    r2 = p.add_run(val)
    r2.font.name = 'Arial'
    r2.font.size = Pt(11)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — Tổng quan kiến trúc dữ liệu
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, '1. Tổng quan kiến trúc dữ liệu')

body(doc, 'Hệ thống quản lý hai loại khách hàng với luồng riêng biệt hội tụ tại bảng trung tâm `subscriptions`.')

arch_diagram = [
    '══════════════════════════════════════════════════════════════════════',
    '                   KHÁCH HÀNG PHỔ THÔNG (INDIVIDUAL)',
    '══════════════════════════════════════════════════════════════════════',
    '',
    '  [plan_templates]──────────[plan_pricing_rules]',
    '        │  customer_segment=INDIVIDUAL',
    '        │',
    '  [retail_plan_schedules]',
    '        │  (lịch áp dụng gói, phải được duyệt)',
    '        │',
    '        ▼',
    '  [approval_requests] ──── [approval_request_steps]',
    '        │  (luồng duyệt đa cấp 1-3 level)',
    '        │',
    '        ▼ (sau khi APPROVED + ACTIVE)',
    '  ╔══════════════════╗',
    '  ║  subscriptions   ║   subscriber_type = INDIVIDUAL',
    '  ║  user_id = ?     ║   (ID khách hàng từ RS Core — không có FK nội bộ)',
    '  ╚══════════════════╝',
    '',
    '══════════════════════════════════════════════════════════════════════',
    '                   KHÁCH HÀNG ĐẠI LÝ (GROUP)',
    '══════════════════════════════════════════════════════════════════════',
    '',
    '  [groups] ──── [group_contacts]',
    '     │      └── [group_members]',
    '     │',
    '  [plan_templates]──────────[plan_pricing_rules]',
    '     │  customer_segment=GROUP',
    '     │',
    '  [group_plan_assignments]',
    '     │  (gói gán cho đại lý, phải được duyệt)',
    '     │',
    '     ▼',
    '  [approval_requests] ──── [approval_request_steps]',
    '     │  (luồng duyệt đa cấp 1-3 level)',
    '     │',
    '     ▼ (sau khi APPROVED + ACTIVE)',
    '  ╔══════════════════╗',
    '  ║  subscriptions   ║   subscriber_type = GROUP',
    '  ║  group_id = ?    ║   (FK trỏ về bảng groups)',
    '  ╚══════════════════╝',
    '',
    '══════════════════════════════════════════════════════════════════════',
    '              SAU KHI SUBSCRIPTION ĐƯỢC TẠO (cả hai loại)',
    '══════════════════════════════════════════════════════════════════════',
    '',
    '  [subscriptions]',
    '        │',
    '        ├──► [certificate_provisioning_records]   (cấp phát chứng thư số)',
    '        │          │',
    '        │          └──► [certificate_usage_records]  (ghi nhận lần ký)',
    '        │',
    '        ├──► [payment_records]                    (ghi nhận thanh toán)',
    '        │',
    '        └──► [usage_aggregates]                   (tổng hợp usage theo kỳ)',
    '                   │',
    '              (chỉ GROUP)',
    '                   └──► [settlement_statements]   (sao kê thanh toán đại lý)',
    '                              │',
    '                              └──► [payment_records] (thanh toán settlement)',
]
add_code_block(doc, arch_diagram)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — Mô tả chi tiết từng bảng
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, '2. Mô tả chi tiết từng bảng')

# 2.1
h2(doc, '2.1 Nhóm bảng định nghĩa gói cước')

h3(doc, 'plan_templates — Mẫu gói cước')
body(doc, 'Lưu trữ định nghĩa các gói dịch vụ. Mỗi gói áp dụng cho một loại khách hàng cụ thể.')

add_md_table(doc,
    ['Cột', 'Kiểu', 'Mô tả'],
    [
        ['`plan_template_id`', 'BIGINT PK', 'Khóa chính'],
        ['`plan_code`', 'VARCHAR UNIQUE', 'Mã gói, ví dụ: `PLN_INDIVIDUAL_001`'],
        ['`plan_name`', 'VARCHAR', 'Tên gói hiển thị'],
        ['`customer_segment`', 'ENUM', '`INDIVIDUAL` hoặc `GROUP` — phân biệt loại KH'],
        ['`template_scope`', 'ENUM', '`PUBLIC` (hiển thị), `PARTNER_PRIVATE`, `SYSTEM`'],
        ['`status`', 'ENUM', '`DRAFT` → `AVAILABLE` → `INACTIVE` → `ARCHIVED`'],
        ['`effective_from` / `effective_to`', 'DATE', 'Khoảng thời gian gói hợp lệ'],
        ['`allow_bulk_signing`', 'BOOL', 'Cho phép ký hàng loạt'],
        ['`allow_api_access`', 'BOOL', 'Cho phép truy cập qua API'],
        ['`cloned_from_template_id`', 'FK self', 'Nếu gói này được clone từ gói khác'],
        ['`version_no`', 'INT', 'Phiên bản gói'],
    ],
    [4.5, 3.5, 9]
)

body(doc, '**Quan hệ:**')
bullet(doc, '1 template → N `plan_pricing_rules` (các mức giá khác nhau)')
bullet(doc, '1 template → N `retail_plan_schedules` (lịch áp dụng cho KH phổ thông)')
bullet(doc, '1 template → N `group_plan_assignments` (gán cho đại lý)')
bullet(doc, '1 template → N `subscriptions` (subscription thực tế)')

divider(doc)

h3(doc, 'plan_pricing_rules — Quy tắc giá')
body(doc, 'Một template có nhiều rule giá, phân theo bậc số lượng hoặc loại chứng thư. **Quan trọng:** rule này được tham chiếu bởi `subscriptions` để tính tiền lúc settlement — không chỉ là audit.')

add_md_table(doc,
    ['Cột', 'Kiểu', 'Mô tả'],
    [
        ['`plan_pricing_rule_id`', 'BIGINT PK', 'Khóa chính'],
        ['`plan_template_id`', 'FK', 'Thuộc template nào'],
        ['`subject_type`', 'ENUM', '`INDIVIDUAL`, `ORGANIZATION`, `INDIVIDUAL_OF_ORG`'],
        ['`certificate_validity_value`', 'INT', 'Thời hạn chứng thư (số)'],
        ['`certificate_validity_unit`', 'ENUM', '`DAY`, `MONTH`, `YEAR`'],
        ['`pricing_metric`', 'ENUM', '`CERTIFICATE_COUNT` (theo số chứng thư) hoặc `SIGNING_COUNT` (theo số lần ký)'],
        ['`range_min` / `range_max`', 'INT', 'Bậc số lượng tối thiểu / tối đa'],
        ['`unit_price`', 'DECIMAL', 'Đơn giá'],
        ['`currency`', 'CHAR(3)', 'Tiền tệ, mặc định `VND`'],
        ['`quota_total`', 'INT', 'Tổng hạn mức (null = không giới hạn)'],
        ['`sort_order`', 'INT', 'Thứ tự hiển thị'],
        ['`is_active`', 'BOOL', 'Rule có đang hiệu lực không'],
    ],
    [4.5, 3.5, 9]
)

add_note(doc, '**Lưu ý thiết kế:** `unit_price` và `currency` **không được copy** vào `subscriptions` — lúc tính settlement, hệ thống đọc trực tiếp từ bảng này qua FK `pricing_rule_id`. Điều này có rủi ro nếu rule bị sửa sau khi subscription tồn tại.')

# 2.2
h2(doc, '2.2 Nhóm bảng khách hàng phổ thông')

h3(doc, 'retail_plan_schedules — Lịch áp dụng gói phổ thông')
body(doc, 'Mỗi `plan_template` có thể có nhiều lịch áp dụng theo thời gian khác nhau. Lịch phải qua luồng duyệt trước khi ACTIVE.')

add_md_table(doc,
    ['Cột', 'Kiểu', 'Mô tả'],
    [
        ['`retail_plan_schedule_id`', 'BIGINT PK', 'Khóa chính'],
        ['`plan_template_id`', 'FK', 'Template gói tương ứng'],
        ['`schedule_status`', 'ENUM', '`AVAILABLE` → `REQUESTED` → `APPROVED` → `ACTIVE` → `INACTIVE`'],
        ['`apply_from` / `apply_to`', 'DATE', 'Khoảng thời gian áp dụng'],
        ['`requested_by` / `requested_at`', 'VARCHAR / DATETIME', 'Ai yêu cầu và khi nào'],
        ['`approved_by` / `approved_at`', 'VARCHAR / DATETIME', 'Ai duyệt và khi nào'],
    ],
    [4.5, 3.5, 9]
)

body(doc, '**Quan hệ:**')
bullet(doc, 'N:1 với `plan_templates`')
bullet(doc, '1:N với `subscriptions` (mỗi khách hàng đăng ký → 1 subscription liên kết lịch này)')
bullet(doc, '1:N với `assignment_audits` (lịch sử thay đổi trạng thái)')

# 2.3
h2(doc, '2.3 Nhóm bảng khách hàng đại lý')

h3(doc, 'groups — Đại lý')
body(doc, 'Thông tin tổ chức của đại lý (khách hàng doanh nghiệp).')

add_md_table(doc,
    ['Cột', 'Kiểu', 'Mô tả'],
    [
        ['`group_id`', 'BIGINT PK', 'Khóa chính'],
        ['`group_code`', 'VARCHAR UNIQUE', 'Mã đại lý duy nhất'],
        ['`group_name`', 'VARCHAR', 'Tên đại lý'],
        ['`username` / `password`', 'VARCHAR', 'Tài khoản đăng nhập của đại lý vào hệ thống'],
        ['`ref_contract_no`', 'VARCHAR', 'Số hợp đồng tham chiếu'],
        ['`status`', 'ENUM', '`ACTIVE` hoặc `INACTIVE`'],
        ['`owner_user_id`', 'FK → `user_accounts`', 'Nhân viên kinh doanh phụ trách'],
    ],
    [4.5, 3.5, 9]
)

body(doc, '**Quan hệ:**')
bullet(doc, '1:N với `group_contacts`')
bullet(doc, '1:N với `group_members`')
bullet(doc, '1:N với `group_plan_assignments`')
bullet(doc, '1:N với `settlement_statements`')
bullet(doc, '1:N với `subscriptions` (qua group_id)')

divider(doc)

h3(doc, 'group_contacts — Đầu mối liên hệ đại lý')
body(doc, 'Lưu các đầu mối liên hệ của đại lý theo từng mục đích: hợp đồng, thanh toán, kỹ thuật, phụ trách chính.')

add_md_table(doc,
    ['Cột', 'Kiểu', 'Mô tả'],
    [
        ['`group_contact_id`', 'BIGINT PK', 'Khóa chính'],
        ['`group_id`', 'FK → `groups`', 'Đại lý'],
        ['`contact_type`', 'ENUM', '`CONTRACT`, `BILLING`, `SUPPORT`, `PIC`'],
        ['`email`', 'VARCHAR', 'Email'],
        ['`full_name`', 'VARCHAR', 'Họ tên'],
        ['`phone`', 'VARCHAR', 'Số điện thoại'],
        ['`is_primary`', 'BOOL', 'Liên hệ chính trong loại này'],
        ['`receive_usage_alert`', 'BOOL', 'Nhận cảnh báo khi gần hết quota'],
        ['`is_active`', 'BOOL', 'Còn hiệu lực không'],
    ],
    [4.5, 3.5, 9]
)

divider(doc)

h3(doc, 'group_members — Thành viên đại lý')
body(doc, 'Danh sách người dùng (user_id từ hệ thống ngoài) thuộc đại lý và vai trò của họ.')

add_md_table(doc,
    ['Cột', 'Kiểu', 'Mô tả'],
    [
        ['`id`', 'BIGINT PK', 'Khóa chính'],
        ['`group_id`', 'FK → `groups`', 'Đại lý'],
        ['`user_id`', 'BIGINT (không FK)', 'ID người dùng từ hệ thống RS Core'],
        ['`role`', 'ENUM', '`OPERATOR` (quản trị viên đại lý) hoặc `MEMBER`'],
        ['`joined_at`', 'DATETIME', 'Thời điểm tham gia'],
        ['`added_by`', 'VARCHAR', 'Ai thêm thành viên này'],
        ['`member_start_date` / `member_end_date`', 'DATE', 'Khoảng thời gian là thành viên'],
        ['`source_assignment_id`', 'FK → `group_plan_assignments`', 'Thành viên được thêm từ gói nào'],
    ],
    [4.5, 3.5, 9]
)

divider(doc)

h3(doc, 'group_plan_assignments — Gói gán cho đại lý')
body(doc, 'Mỗi bản ghi là một lần gán gói cước cho một đại lý, có khoảng thời gian áp dụng. Phải qua luồng duyệt trước khi ACTIVE.')

add_md_table(doc,
    ['Cột', 'Kiểu', 'Mô tả'],
    [
        ['`group_plan_assignment_id`', 'BIGINT PK', 'Khóa chính'],
        ['`group_id`', 'FK → `groups`', 'Đại lý'],
        ['`plan_template_id`', 'FK → `plan_templates`', 'Gói cước áp dụng'],
        ['`assignment_status`', 'ENUM', '`REQUESTED` → `APPROVED` → `ACTIVE` → `STOPPED` / `EXPIRED`'],
        ['`requested_by` / `requested_at`', '—', 'Ai tạo yêu cầu'],
        ['`approved_by` / `approved_at`', '—', 'Ai duyệt'],
        ['`rejected_by` / `rejected_at`', '—', 'Ai từ chối (nếu có)'],
        ['`apply_from` / `apply_to`', 'DATE', 'Khoảng thời gian hiệu lực'],
        ['`activated_at`', 'DATETIME', 'Thời điểm kích hoạt'],
        ['`stopped_at` / `stop_reason`', '—', 'Dừng thủ công kèm lý do'],
    ],
    [4.5, 3.5, 9]
)

body(doc, '**Quan hệ:**')
bullet(doc, 'N:1 với `groups` và `plan_templates`')
bullet(doc, '1:N với `assignment_audits`')
bullet(doc, '1:N với `subscriptions`')
bullet(doc, '1:N với `certificate_provisioning_records`')

# 2.4
h2(doc, '2.4 Bảng duyệt chung (dùng cho cả hai loại KH)')

h3(doc, 'approval_requests — Yêu cầu phê duyệt')
body(doc, 'Mỗi lần tạo lịch áp dụng (retail) hoặc gán gói (group) đều sinh ra 1 approval request và phải đi qua đây trước khi thực thi.')

add_md_table(doc,
    ['Cột', 'Kiểu', 'Mô tả'],
    [
        ['`id`', 'BIGINT PK', 'Khóa chính'],
        ['`request_type`', 'VARCHAR', 'Loại yêu cầu: `REQUEST_GROUP_PLAN_ASSIGNMENT`, `REQUEST_RETAIL_PLAN_SCHEDULE`, ...'],
        ['`status`', 'VARCHAR', '`DRAFT` → `IN_APPROVAL` → `APPROVED` / `REJECTED` / `NEED_REVISION`'],
        ['`customer_segment`', 'VARCHAR', '`INDIVIDUAL` hoặc `GROUP`'],
        ['`requested_by`', 'VARCHAR', 'Người tạo yêu cầu'],
        ['`entity_type`', 'VARCHAR', '`RETAIL_PLAN_SCHEDULE` hoặc `GROUP_PLAN_ASSIGNMENT`'],
        ['`entity_id`', 'VARCHAR', 'ID của schedule hoặc assignment tương ứng'],
        ['`request_payload`', 'TEXT', 'Snapshot dữ liệu lúc tạo yêu cầu (JSON)'],
        ['`contract_value`', 'DECIMAL', 'Giá trị hợp đồng — dùng để tính số cấp phải duyệt'],
        ['`total_levels`', 'INT', 'Tổng số cấp duyệt (1, 2 hoặc 3)'],
        ['`current_level`', 'INT', 'Đang ở cấp duyệt nào'],
    ],
    [4.5, 3.5, 9]
)
body(doc, '**Logic số cấp duyệt** được tra từ bảng `approval_level_configs` theo `customer_segment` và ngưỡng `contract_value`.')

divider(doc)

h3(doc, 'approval_request_steps — Từng bước duyệt')
body(doc, 'Mỗi approval request có 1-3 steps tương ứng 1-3 cấp duyệt.')

add_md_table(doc,
    ['Cột', 'Kiểu', 'Mô tả'],
    [
        ['`id`', 'BIGINT PK', 'Khóa chính'],
        ['`approval_request_id`', 'FK', 'Thuộc request nào'],
        ['`step_level`', 'INT', 'Cấp duyệt: 1, 2 hoặc 3'],
        ['`required_approval_level`', 'ENUM', '`LEVEL_1` (Trưởng phòng), `LEVEL_2` (Giám đốc), `LEVEL_3` (CFO)'],
        ['`status`', 'ENUM', '`PENDING` → `APPROVED` / `REJECTED` / `SKIPPED`'],
        ['`decided_by`', 'VARCHAR', 'Ai duyệt bước này'],
        ['`comment`', 'TEXT', 'Ghi chú khi duyệt'],
        ['`decided_at`', 'DATETIME', 'Thời điểm quyết định'],
    ],
    [4.5, 3.5, 9]
)

# 2.5
h2(doc, '2.5 Bảng trung tâm — subscriptions')
body(doc, 'Đây là bảng **trung tâm** — mỗi bản ghi là một hợp đồng dịch vụ đang hoặc đã tồn tại của một khách hàng.')

add_md_table(doc,
    ['Cột', 'Kiểu', 'Mô tả'],
    [
        ['`subscription_id`', 'BIGINT PK', 'Khóa chính'],
        ['`subscriber_type`', 'ENUM', '`INDIVIDUAL` hoặc `GROUP`'],
        ['`user_id`', 'BIGINT (không FK)', 'ID khách hàng từ RS Core — chỉ điền khi `INDIVIDUAL`'],
        ['`group_id`', 'FK → `groups`', 'Đại lý — chỉ điền khi `GROUP`'],
        ['`plan_template_id`', 'FK → `plan_templates`', 'Gói cước đã đăng ký'],
        ['`pricing_rule_id`', 'FK → `plan_pricing_rules`', 'Rule giá được áp dụng — dùng để tính tiền lúc settlement'],
        ['`group_plan_assignment_id`', 'FK', 'Liên kết về gán gói của đại lý (nếu GROUP)'],
        ['`retail_plan_schedule_id`', 'FK', 'Liên kết về lịch áp dụng (nếu INDIVIDUAL)'],
        ['`source_type`', 'ENUM', '`RETAIL_PURCHASE`, `GROUP_ASSIGNMENT`, `MANUAL`'],
        ['`status`', 'ENUM', '`PENDING` → `ACTIVE` → `EXPIRED` / `CANCELLED` / `SUSPENDED`'],
        ['`start_date` / `end_date`', 'DATE', 'Thời gian hiệu lực'],
        ['`signing_quota_total`', 'INT', 'Hạn mức chữ ký — **copy từ pricing_rule.quota_total** lúc tạo, sau đó là field độc lập'],
        ['`signing_quota_used`', 'INT', 'Đã dùng bao nhiêu'],
        ['`activated_by`', 'VARCHAR', 'Ai kích hoạt'],
        ['`payment_reference`', 'VARCHAR', 'Mã tham chiếu thanh toán'],
    ],
    [4.5, 3.5, 9]
)

add_note(doc, '**Tại sao có cả pricing_rule_id lẫn signing_quota_total?** signing_quota_total được copy từ pricing_rule.quota_total tại thời điểm tạo subscription → cố định, không thay đổi dù rule bị sửa. Còn pricing_rule_id (FK) được dùng để đọc unit_price và currency live khi tính tiền settlement. Đây là rủi ro tiềm ẩn nếu đơn giá bị sửa sau khi subscription đã tồn tại.')

add_note(doc, '**Tại sao user_id không có FK?** user_accounts chứa nhân viên nội bộ (Sale, Manager, Approver). Khách hàng phổ thông là người dùng bên ngoài, quản lý bởi hệ thống RS Core. user_id trong subscriptions là ID từ RS Core — hệ thống billing chỉ lưu để tham chiếu.')

# 2.6
h2(doc, '2.6 Bảng sau subscription')

h3(doc, 'certificate_provisioning_records — Cấp phát chứng thư số')
body(doc, 'Mỗi khi subscription ACTIVE, hệ thống gọi RS Core để cấp chứng thư. Mỗi lần gọi = 1 bản ghi.')

add_md_table(doc,
    ['Cột', 'Kiểu', 'Mô tả'],
    [
        ['`id`', 'BIGINT PK', 'Khóa chính'],
        ['`subscription_id`', 'FK (no constraint)', 'Subscription tương ứng'],
        ['`group_plan_assignment_id`', 'FK (no constraint)', 'Assignment nếu là GROUP'],
        ['`pricing_rule_id`', 'FK (no constraint)', 'Rule giá áp dụng cho chứng thư này'],
        ['`user_id`', 'BIGINT', 'Người nhận chứng thư (từ RS Core)'],
        ['`request_id`', 'VARCHAR UNIQUE', 'ID request gửi lên RS Core — dùng để chống duplicate'],
        ['`status`', 'ENUM', '`PENDING` → `COMPLETED` / `FAILED` / `FAILED_PERMANENT`'],
        ['`cert_type`', 'TINYINT', '`1`=Cá nhân, `2`=Cá nhân thuộc tổ chức, `3`=Tổ chức'],
        ['`certificate_id`', 'VARCHAR', 'ID chứng thư trả về từ RS Core'],
        ['`key_id`', 'VARCHAR', 'Key ID đi kèm chứng thư'],
        ['`issued_at` / `expires_at`', 'DATETIME', 'Thời gian cấp và hết hạn chứng thư'],
        ['`retry_count`', 'INT', 'Số lần thử lại khi thất bại'],
        ['`usage_count`', 'INT', 'Số lần chứng thư này được dùng'],
        ['`failure_reason`', 'TEXT', 'Lý do thất bại nếu có'],
    ],
    [4.5, 3.5, 9]
)
add_note(doc, '**Tại sao FK là NO_CONSTRAINT?** Cho phép xoá subscription mà không ảnh hưởng đến lịch sử provisioning. Đây là dữ liệu lịch sử — cần giữ lại ngay cả khi subscription bị xoá.')

divider(doc)

h3(doc, 'certificate_usage_records — Lịch sử sử dụng chứng thư')
body(doc, 'Mỗi lần khách hàng ký tài liệu hoặc tạo/gia hạn chứng thư = 1 bản ghi.')

add_md_table(doc,
    ['Cột', 'Kiểu', 'Mô tả'],
    [
        ['`id`', 'BIGINT PK', 'Khóa chính'],
        ['`certificate_id`', 'VARCHAR', 'ID chứng thư đã dùng'],
        ['`user_id`', 'BIGINT', 'Người dùng thực hiện hành động'],
        ['`subscription_id`', 'FK (no constraint)', 'Subscription tương ứng'],
        ['`group_plan_assignment_id`', 'FK (no constraint)', 'Assignment (nếu GROUP)'],
        ['`usage_type`', 'ENUM', '`SIGNING`, `CERTIFICATE_CREATED`, `CERTIFICATE_RENEWED`, `CERTIFICATE_REVOKED`'],
        ['`quantity`', 'INT', 'Số lượng (thường = 1)'],
        ['`used_at`', 'DATETIME', 'Thời điểm sử dụng'],
    ],
    [4.5, 3.5, 9]
)
add_note(doc, '**Vai trò trong billing đại lý:** Khi tính settlement, hệ thống quét toàn bộ certificate_usage_records trong kỳ, đọc pricing_rule qua subscription để tính totalAmount.')

divider(doc)

h3(doc, 'payment_records — Bản ghi thanh toán')
body(doc, 'Ghi nhận các giao dịch thanh toán, liên kết linh hoạt với subscription hoặc settlement.')

add_md_table(doc,
    ['Cột', 'Kiểu', 'Mô tả'],
    [
        ['`payment_id`', 'BIGINT PK', 'Khóa chính'],
        ['`subscription_id`', 'FK → `subscriptions`', 'Thanh toán cho subscription cụ thể'],
        ['`group_plan_assignment_id`', 'FK', 'Thanh toán cho gói đại lý'],
        ['`settlement_statement_id`', 'FK', 'Thanh toán theo sao kê'],
        ['`external_reference`', 'VARCHAR UNIQUE', 'Mã giao dịch từ cổng thanh toán'],
        ['`amount`', 'DECIMAL', 'Số tiền'],
        ['`currency`', 'CHAR(3)', 'Tiền tệ'],
        ['`payment_status`', 'ENUM', '`SUCCESS`, `FAILED`, `REFUNDED`'],
        ['`payment_scope`', 'ENUM', '`SUBSCRIPTION` (cá nhân trả thẳng), `GROUP_ASSIGNMENT`, `SETTLEMENT` (đại lý thanh toán theo kỳ)'],
        ['`payment_method`', 'VARCHAR', 'Phương thức: chuyển khoản, thẻ, ...'],
        ['`paid_at`', 'DATETIME', 'Thời điểm thanh toán'],
        ['`raw_payload`', 'JSON', 'Dữ liệu thô từ cổng thanh toán'],
    ],
    [4.5, 3.5, 9]
)

divider(doc)

h3(doc, 'usage_aggregates — Tổng hợp usage theo kỳ')
body(doc, 'Bảng rollup — tổng hợp usage của group hoặc individual theo kỳ (ngày/tháng) để phục vụ báo cáo và tính settlement nhanh.')

add_md_table(doc,
    ['Cột', 'Mô tả'],
    [
        ['`aggregate_scope`', '`GROUP`, `GROUP_ASSIGNMENT`, `USER`, `RETAIL_PLAN`'],
        ['`scope_id`', 'ID tương ứng với scope'],
        ['`period_type`', '`DAY` hoặc `MONTH`'],
        ['`period_key`', 'Dạng `YYYYMMDD-YYYYMMDD`'],
        ['`certificates_created`', 'Số chứng thư đã tạo'],
        ['`signing_used`', 'Số lần ký'],
        ['`amount_due`', 'Số tiền phát sinh'],
        ['`currency`', 'Tiền tệ'],
    ],
    [5, 12]
)

divider(doc)

h3(doc, 'settlement_statements — Sao kê thanh toán (chỉ đại lý)')
body(doc, 'Tóm tắt chi phí của một đại lý trong một kỳ. Chỉ áp dụng cho `GROUP`.')

add_md_table(doc,
    ['Cột', 'Kiểu', 'Mô tả'],
    [
        ['`settlement_statement_id`', 'BIGINT PK', 'Khóa chính'],
        ['`group_id`', 'FK → `groups`', 'Đại lý'],
        ['`from_date` / `to_date`', 'DATE', 'Kỳ sao kê'],
        ['`status`', 'ENUM', '`DRAFT` → `FINALIZED` → `EXPORTED`'],
        ['`total_certificates`', 'INT', 'Tổng số chứng thư cấp trong kỳ'],
        ['`total_signings`', 'INT', 'Tổng số lần ký trong kỳ'],
        ['`total_amount`', 'DECIMAL', 'Tổng tiền phát sinh'],
        ['`currency`', 'CHAR(3)', 'Tiền tệ'],
        ['`generated_by` / `generated_at`', '—', 'Ai tạo sao kê và khi nào'],
    ],
    [4.5, 3.5, 9]
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — Mối quan hệ
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, '3. Mối quan hệ giữa các bảng')

rel_lines = [
    'plan_templates (1) ─────────── (N) plan_pricing_rules',
    'plan_templates (1) ─────────── (N) retail_plan_schedules',
    'plan_templates (1) ─────────── (N) group_plan_assignments',
    'plan_templates (1) ─────────── (N) subscriptions',
    '',
    'retail_plan_schedules (1) ──── (N) subscriptions',
    'retail_plan_schedules (1) ──── (N) assignment_audits',
    '',
    'groups (1) ─────────────────── (N) group_contacts',
    'groups (1) ─────────────────── (N) group_members',
    'groups (1) ─────────────────── (N) group_plan_assignments',
    'groups (1) ─────────────────── (N) subscriptions',
    'groups (1) ─────────────────── (N) settlement_statements',
    '',
    'group_plan_assignments (1) ──── (N) subscriptions',
    'group_plan_assignments (1) ──── (N) assignment_audits',
    'group_plan_assignments (1) ──── (N) certificate_provisioning_records',
    'group_plan_assignments (1) ──── (N) group_members (via source_assignment_id)',
    '',
    'approval_requests (1) ──────── (N) approval_request_steps',
    '',
    'subscriptions (1) ──────────── (N) certificate_provisioning_records',
    'subscriptions (1) ──────────── (N) certificate_usage_records',
    'subscriptions (1) ──────────── (N) payment_records',
    '',
    'settlement_statements (1) ──── (N) payment_records',
    '',
    'plan_pricing_rules (1) ──────── (N) subscriptions          (snapshot giá)',
    'plan_pricing_rules (1) ──────── (N) certificate_provisioning_records',
]
add_code_block(doc, rel_lines)

h2(doc, 'Sơ đồ ERD rút gọn')

erd = [
    '┌─────────────────────┐         ┌──────────────────────────┐',
    '│    plan_templates    │─────1:N─│    plan_pricing_rules    │',
    '└──────────┬──────────┘         └────────────┬─────────────┘',
    '           │                                 │',
    '      ┌────┴──────────────────┐              │',
    '      │                       │              │',
    ' 1:N  ▼                  1:N  ▼              │',
    '┌──────────────────┐   ┌──────────────────────┐│',
    '│retail_plan_      │   │group_plan_assignments ││',
    '│schedules         │   └────────┬──────────────┘│',
    '└────────┬─────────┘            │               │',
    '         │                      │               │',
    '         └──────────┬───────────┘               │',
    '                    │  1:N                      │',
    '                    ▼                           │',
    '        ┌───────────────────────────────┐       │',
    '        │         subscriptions         │◄──────┘ (FK pricing_rule_id)',
    '        │  subscriber_type: IND | GROUP │',
    '        └──────────┬────────────────────┘',
    '                   │',
    '        ┌──────────┼──────────────────────┐',
    '        │ 1:N      │ 1:N                  │ 1:N',
    '        ▼          ▼                      ▼',
    '┌───────────┐ ┌──────────────┐    ┌───────────────┐',
    '│certificate│ │certificate_  │    │payment_records│',
    '│_provision │ │usage_records │    └───────────────┘',
    '│_records   │ └──────┬───────┘',
    '└───────────┘        │ (tính settlement)',
    '                     ▼',
    '             ┌──────────────────┐',
    '             │usage_aggregates  │',
    '             └────────┬─────────┘',
    '                      │ (chỉ GROUP)',
    '                      ▼',
    '             ┌──────────────────────┐',
    '             │settlement_statements │',
    '             └──────────────────────┘',
]
add_code_block(doc, erd)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — Luồng Individual
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, '4. Luồng nghiệp vụ — Khách hàng Phổ thông (INDIVIDUAL)')

steps_individual = [
    ('Bước 1: Tạo gói cước',
     [('Ai', 'Admin'), ('API', 'POST /api/v1/individual/plan-configs')],
     ['INSERT plan_templates          (customer_segment = INDIVIDUAL, status = AVAILABLE)',
      'INSERT plan_pricing_rules (N)  (các mức giá theo bậc số lượng)']),
    ('Bước 2: Tạo lịch áp dụng và gửi duyệt',
     [('Ai', 'Admin'), ('API', 'POST /api/v1/individual/plan-configs/{id}/request-apply')],
     ['INSERT retail_plan_schedules   (schedule_status = REQUESTED)',
      'INSERT assignment_audits       (action = REQUEST)',
      'INSERT approval_requests       (status = IN_APPROVAL, customer_segment = INDIVIDUAL,',
      '                                entity_type = RETAIL_PLAN_SCHEDULE,',
      '                                contract_value = max(unit_price) từ pricing_rules)',
      'INSERT approval_request_steps  (1-3 steps tùy ngưỡng contract_value)',
      '→ gửi email cho APPROVAL_L1']),
    ('Bước 3: Phê duyệt đa cấp',
     [('Ai', 'Approver (Level 1 / 2 / 3)'), ('API', 'POST /api/v1/approval-requests/{id}/approve')],
     ['UPDATE approval_request_steps  (status = APPROVED cho step hiện tại)',
      'UPDATE approval_requests       (current_level++)',
      '',
      'Nếu đủ tất cả level:',
      'UPDATE approval_requests       (status = APPROVED)',
      'UPDATE retail_plan_schedules   (schedule_status = APPROVED, approved_by, approved_at)',
      'INSERT assignment_audits       (action = APPROVE)',
      '',
      'Nếu bị từ chối:',
      'UPDATE approval_request_steps  (status = REJECTED)',
      'UPDATE approval_requests       (status = REJECTED)',
      'UPDATE retail_plan_schedules   (schedule_status = INACTIVE)',
      'INSERT assignment_audits       (action = REJECT)']),
    ('Bước 4: Kích hoạt lịch',
     [('Khi nào', 'Scheduler hàng ngày hoặc admin thủ công'),
      ('Điều kiện', 'schedule_status = APPROVED và apply_from <= today')],
     ['UPDATE retail_plan_schedules   (schedule_status = ACTIVE)',
      'INSERT assignment_audits       (action = ACTIVATE)']),
    ('Bước 5: Khách hàng đăng ký subscription',
     [('Ai', 'Khách hàng tự đăng ký hoặc Sale tạo thủ công'),
      ('API', 'POST /api/v1/commercial-flows/retail-plan-schedules/{scheduleId}/execute')],
     ['1. Kiểm tra schedule_status = ACTIVE',
      '2. Resolve pricing rule (ưu tiên rule truyền vào, nếu không lấy rule mặc định)',
      '3. INSERT subscriptions:',
      '      subscriber_type = INDIVIDUAL',
      '      user_id         = <ID khách hàng từ RS Core>',
      '      plan_template_id',
      '      pricing_rule_id',
      '      retail_plan_schedule_id',
      '      source_type     = RETAIL_PURCHASE',
      '      status          = ACTIVE (hoặc PENDING nếu chờ thanh toán)',
      '      start_date, end_date',
      '      signing_quota_total = pricing_rule.quota_total  ← copy vào đây',
      '      signing_quota_used  = 0']),
    ('Bước 6: Cấp phát chứng thư',
     [('Khi nào', 'Ngay sau khi subscription ACTIVE'),
      ('Service', 'CertificateProvisioningService')],
     ['INSERT certificate_provisioning_records:',
      '      subscription_id',
      '      user_id          = <ID khách hàng>',
      '      request_id       = UUID (gửi lên RS Core)',
      '      status           = PENDING',
      '      cert_type        = INDIVIDUAL (1)',
      '→ Gọi RS Core API để tạo chứng thư',
      '→ RS Core trả về certificate_id, key_id, expires_at',
      'UPDATE certificate_provisioning_records:',
      '      status           = COMPLETED',
      '      certificate_id, key_id, issued_at, expires_at']),
    ('Bước 7: Khách hàng sử dụng chứng thư',
     [],
     ['INSERT certificate_usage_records:',
      '      certificate_id',
      '      user_id',
      '      subscription_id',
      '      usage_type       = SIGNING (hoặc CERTIFICATE_CREATED, ...)',
      '      quantity         = 1',
      '      used_at          = NOW()',
      '',
      'UPDATE subscriptions:',
      '      signing_quota_used++   ← trừ quota']),
    ('Bước 8: Hết hạn',
     [('Khi nào', 'Scheduler chạy, phát hiện end_date < today')],
     ['UPDATE subscriptions                  (status = EXPIRED)',
      'UPDATE retail_plan_schedules          (schedule_status = INACTIVE) nếu hết kỳ',
      'INSERT assignment_audits              (action = EXPIRE)']),
]

for title, meta, code in steps_individual:
    h3(doc, title)
    for label, val in meta:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        r1 = p.add_run(label + ': ')
        r1.bold = True; r1.font.name = 'Arial'; r1.font.size = Pt(10.5)
        add_inline(p, val)
    if code:
        add_code_block(doc, code)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — Luồng Group
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, '5. Luồng nghiệp vụ — Khách hàng Đại lý (GROUP)')

steps_group = [
    ('Bước 1: Tạo đại lý',
     [('Ai', 'Admin'), ('API', 'POST /api/v1/groups')],
     ['INSERT groups              (status = ACTIVE)',
      'INSERT group_contacts (N)  (tùy chọn — thêm sau cũng được)']),
    ('Bước 2: Gán gói cước cho đại lý và gửi duyệt',
     [('Ai', 'Admin / Sale'),
      ('API (A)', 'POST /api/v1/groups/{groupId}/plan-assignments  — gán template đã có'),
      ('API (B)', 'POST /api/v1/groups/{groupId}/add-plan  — tạo template mới luôn'),
      ('API (C)', 'POST /api/v1/groups/provision  — tạo group + template + assignment cùng lúc')],
     ['INSERT plan_templates          (customer_segment = GROUP)    — nếu tạo mới',
      'INSERT plan_pricing_rules (N)                                — nếu tạo mới',
      'INSERT group_plan_assignments  (assignment_status = REQUESTED)',
      'INSERT assignment_audits       (action = REQUEST)',
      'INSERT approval_requests       (status = IN_APPROVAL, customer_segment = GROUP,',
      '                                entity_type = GROUP_PLAN_ASSIGNMENT,',
      '                                contract_value = max(unit_price * quota_total))',
      'INSERT approval_request_steps  (1-3 steps)',
      '→ gửi email cho APPROVAL_L1']),
    ('Bước 3: Phê duyệt đa cấp',
     [('Ai', 'Approver (Level 1 / 2 / 3)'),
      ('API', 'POST /api/v1/approval-requests/{id}/approve')],
     ['UPDATE approval_requests       (status = APPROVED)',
      'UPDATE group_plan_assignments  (assignment_status = APPROVED, approved_by, approved_at)',
      'INSERT assignment_audits       (action = APPROVE)']),
    ('Bước 4: Kích hoạt gói',
     [('Ai', 'Admin hoặc Scheduler'),
      ('API', 'POST /api/v1/commercial-flows/group-assignments/{assignmentId}/execute')],
     ['UPDATE group_plan_assignments  (assignment_status = ACTIVE, activated_at = NOW())',
      'INSERT assignment_audits       (action = ACTIVATE)']),
    ('Bước 5: Tạo subscription cho đại lý',
     [('Ghi chú', 'Xảy ra trong cùng request execute ở Bước 4, khi issueSubscription = true')],
     ['1. Kiểm tra chưa có subscription ACTIVE cho assignment này',
      '2. Resolve pricing rule (SIGNING_COUNT được ưu tiên cho đại lý)',
      '3. INSERT subscriptions:',
      '      subscriber_type        = GROUP',
      '      group_id               = <group_id>',
      '      user_id                = NULL         ← không dùng cho GROUP',
      '      plan_template_id',
      '      pricing_rule_id',
      '      group_plan_assignment_id',
      '      source_type            = GROUP_ASSIGNMENT',
      '      status                 = ACTIVE',
      '      start_date, end_date',
      '      signing_quota_total    = pricing_rule.quota_total  ← copy vào đây',
      '      signing_quota_used     = 0']),
    ('Bước 6: Thêm thành viên đại lý',
     [],
     ['INSERT group_members:',
      '      group_id',
      '      user_id                = <ID từ RS Core>',
      '      role                   = OPERATOR / MEMBER',
      '      source_assignment_id   = group_plan_assignment_id']),
    ('Bước 7: Cấp phát chứng thư cho từng thành viên',
     [],
     ['Với mỗi user_id trong group_members:',
      'INSERT certificate_provisioning_records:',
      '      subscription_id',
      '      group_plan_assignment_id',
      '      pricing_rule_id',
      '      user_id',
      '      cert_type              = ORGANIZATION (3) hoặc INDIVIDUAL_OF_ORG (2)',
      '      status                 = PENDING',
      '→ Gọi RS Core API → nhận certificate_id, key_id, expires_at',
      'UPDATE certificate_provisioning_records: status = COMPLETED']),
    ('Bước 8: Thành viên sử dụng chứng thư (ký tài liệu)',
     [],
     ['INSERT certificate_usage_records:',
      '      certificate_id',
      '      user_id',
      '      subscription_id        ← subscription của group',
      '      group_plan_assignment_id',
      '      usage_type             = SIGNING',
      '      quantity               = 1',
      '',
      'UPDATE subscriptions:',
      '      signing_quota_used++']),
    ('Bước 9: Tính sao kê cuối kỳ',
     [('Ai', 'Admin'),
      ('API', 'POST /api/v1/commercial-flows/groups/{groupId}/settlement')],
     ['1. Quét certificate_usage_records trong khoảng from_date → to_date',
      '2. Với mỗi record:',
      '      - Lấy subscription → lấy pricing_rule (LIVE qua FK!)',
      '      - Tính tiền: quantity × pricing_rule.unit_price',
      '3. INSERT settlement_statements:',
      '      group_id',
      '      from_date, to_date',
      '      status               = DRAFT (hoặc FINALIZED nếu finalizeNow=true)',
      '      total_certificates',
      '      total_signings',
      '      total_amount',
      '      currency             (lấy từ pricing_rule.currency)',
      '4. UPSERT usage_aggregates:',
      '      aggregate_scope = GROUP',
      '      scope_id        = group_id',
      '      period_key      = "YYYYMMDD-YYYYMMDD"']),
    ('Bước 10: Hết hạn gói',
     [('Khi nào', 'Scheduler chạy, phát hiện apply_to < today')],
     ['UPDATE group_plan_assignments  (assignment_status = EXPIRED)',
      'UPDATE subscriptions           (status = EXPIRED)',
      'INSERT assignment_audits       (action = EXPIRE)']),
]

for title, meta, code in steps_group:
    h3(doc, title)
    for label, val in meta:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        r1 = p.add_run(label + ': ')
        r1.bold = True; r1.font.name = 'Arial'; r1.font.size = Pt(10.5)
        add_inline(p, val)
    if code:
        add_code_block(doc, code)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — Sau subscription
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, '6. Dữ liệu đi đâu sau khi Subscription được tạo')

flow_after = [
    '                   ┌─────────────────┐',
    '                   │  subscriptions  │ ← điểm bắt đầu',
    '                   └────────┬────────┘',
    '                            │',
    '     ┌─────────────────────┼─────────────────────┐',
    '     │                     │                     │',
    '     ▼                     ▼                     ▼',
    '┌──────────────┐   ┌──────────────────┐   ┌────────────────┐',
    '│ Cấp phát     │   │ Ghi nhận thanh   │   │ Trừ quota      │',
    '│ chứng thư số │   │ toán             │   │ sử dụng        │',
    '│              │   │                  │   │                │',
    '│ certificate_ │   │ payment_records  │   │ signing_quota_ │',
    '│ provisioning │   │ (scope=          │   │ used++         │',
    '│ _records     │   │  SUBSCRIPTION)   │   └────────────────┘',
    '└──────┬───────┘   └──────────────────┘',
    '       │ Khách hàng dùng chứng thư',
    '       ▼',
    '┌──────────────────────┐',
    '│ certificate_usage_   │   ← mỗi lần ký / tạo / gia hạn chứng thư',
    '│ records              │',
    '└────────┬─────────────┘',
    '         │ Rollup định kỳ (scheduler)',
    '         ▼',
    '┌──────────────────────┐',
    '│ usage_aggregates     │   ← tổng hợp theo ngày / tháng',
    '└────────┬─────────────┘',
    '         │ Chỉ với GROUP — cuối kỳ tính sao kê',
    '         ▼',
    '┌──────────────────────┐',
    '│ settlement_          │   DRAFT → FINALIZED → EXPORTED',
    '│ statements           │',
    '└────────┬─────────────┘',
    '         ▼',
    '┌──────────────────────┐',
    '│ payment_records      │   ← đại lý thanh toán theo sao kê',
    '│ (scope = SETTLEMENT) │',
    '└──────────────────────┘',
]
add_code_block(doc, flow_after)

h2(doc, 'Điểm khác biệt then chốt giữa hai loại KH sau subscription')

add_md_table(doc,
    ['Hành động', 'Khách hàng Phổ thông (INDIVIDUAL)', 'Khách hàng Đại lý (GROUP)'],
    [
        ['Ai nhận chứng thư', 'Chính khách hàng (`user_id`)', 'Từng thành viên trong `group_members`'],
        ['Loại chứng thư', '`cert_type = 1` (Cá nhân)', '`cert_type = 2` hoặc `3`'],
        ['Thanh toán', 'Trả thẳng khi đăng ký (`SUBSCRIPTION`)', 'Thanh toán theo sao kê cuối kỳ (`SETTLEMENT`)'],
        ['Sao kê', 'Không có', 'Có `settlement_statements` mỗi kỳ'],
        ['Cách tính tiền', 'Trả trước theo gói', 'Tính theo thực tế sử dụng (ký / chứng thư)'],
        ['`pricing_metric`', '`CERTIFICATE_COUNT`', '`SIGNING_COUNT`'],
        ['Quota tracking', '`signing_quota_total` / `_used` trong subscription', 'Tổng hợp qua `usage_aggregates`'],
    ],
    [4.5, 6.5, 6]
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — Vòng đời trạng thái
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, '7. Vòng đời trạng thái')

states = [
    ('retail_plan_schedules.schedule_status',
     ['AVAILABLE → REQUESTED → APPROVED → ACTIVE → INACTIVE',
      '                │                              ▲',
      '                └──────── REJECTED ────────────┘ (cũng là INACTIVE)']),
    ('group_plan_assignments.assignment_status',
     ['REQUESTED → APPROVED → ACTIVE → STOPPED',
      '    │                    └────→ EXPIRED',
      '    └────→ REJECTED']),
    ('subscriptions.status',
     ['PENDING → ACTIVE → EXPIRED',
      '              ├──→ CANCELLED',
      '              └──→ SUSPENDED']),
    ('certificate_provisioning_records.status',
     ['PENDING → COMPLETED',
      '    └───→ FAILED (retry_count++) → FAILED_PERMANENT']),
    ('approval_requests.status',
     ['DRAFT → IN_APPROVAL → APPROVED',
      '            ├────────→ REJECTED',
      '            └────────→ NEED_REVISION → IN_APPROVAL (resubmit)']),
    ('settlement_statements.status',
     ['DRAFT → FINALIZED → EXPORTED']),
]

for table_name, diagram in states:
    h3(doc, table_name)
    add_code_block(doc, diagram)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 8 — Tóm tắt
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, '8. Bảng tóm tắt nghiệp vụ vs bảng DB')

h2(doc, 'Khách hàng Phổ thông')

add_md_table(doc,
    ['Nghiệp vụ', 'INSERT vào bảng', 'UPDATE bảng'],
    [
        ['Tạo gói cước', '`plan_templates`, `plan_pricing_rules`', '—'],
        ['Gửi duyệt lịch áp dụng', '`retail_plan_schedules`, `assignment_audits`, `approval_requests`, `approval_request_steps`', '—'],
        ['Duyệt / Từ chối', '`assignment_audits`', '`retail_plan_schedules`, `approval_requests`, `approval_request_steps`'],
        ['Kích hoạt lịch', '`assignment_audits`', '`retail_plan_schedules`'],
        ['KH đăng ký subscription', '`subscriptions`', '—'],
        ['Cấp phát chứng thư', '`certificate_provisioning_records`', '`certificate_provisioning_records`'],
        ['KH dùng chứng thư', '`certificate_usage_records`', '`subscriptions` (quota_used++)'],
        ['Hết hạn', '`assignment_audits`', '`subscriptions`, `retail_plan_schedules`'],
    ],
    [4.5, 7.5, 5]
)

h2(doc, 'Khách hàng Đại lý')

add_md_table(doc,
    ['Nghiệp vụ', 'INSERT vào bảng', 'UPDATE bảng'],
    [
        ['Tạo đại lý', '`groups`, `group_contacts`', '—'],
        ['Thêm thành viên', '`group_members`', '—'],
        ['Tạo gói + gán', '`plan_templates`, `plan_pricing_rules`, `group_plan_assignments`, `assignment_audits`, `approval_requests`, `approval_request_steps`', '—'],
        ['Duyệt / Từ chối', '`assignment_audits`', '`group_plan_assignments`, `approval_requests`, `approval_request_steps`'],
        ['Kích hoạt gói', '`assignment_audits`, `subscriptions`', '`group_plan_assignments`'],
        ['Cấp chứng thư cho thành viên', '`certificate_provisioning_records`', '`certificate_provisioning_records`'],
        ['Thành viên ký tài liệu', '`certificate_usage_records`', '`subscriptions` (quota_used++)'],
        ['Tính sao kê', '`settlement_statements`, `usage_aggregates`', '—'],
        ['Đại lý thanh toán sao kê', '`payment_records`', '—'],
        ['Hết hạn gói', '`assignment_audits`', '`group_plan_assignments`, `subscriptions`'],
    ],
    [4.5, 7.5, 5]
)

# ─── Save ─────────────────────────────────────────────────────────────────────
output_path = r'E:\PhuongCM\billing-CA\docs\SUBSCRIPTION_DATA_FLOW.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
