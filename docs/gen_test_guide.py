# -*- coding: utf-8 -*-
"""Generate API test guide Word document for billing-CA subscription APIs."""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page setup A4 ──────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
    setattr(section, attr, Cm(2))

# ── Colors ──────────────────────────────────────────────────────────────────
C_H1      = RGBColor(0x1F, 0x38, 0x64)
C_H2      = RGBColor(0x2E, 0x54, 0x96)
C_H3      = RGBColor(0x21, 0x63, 0xAB)
C_GREEN   = RGBColor(0x1A, 0x6B, 0x2E)  # GET
C_BLUE    = RGBColor(0x0B, 0x4F, 0x8A)  # POST
C_ORANGE  = RGBColor(0xB4, 0x5A, 0x00)  # PATCH
C_BG_GRN  = "E2EFDA"
C_BG_BLU  = "DDEEFF"
C_BG_ORG  = "FFF2CC"
C_BG_CODE = "F2F2F2"
C_BG_HDR  = "D5E8F0"
C_BG_NOTE = "FFF8DC"
C_WHITE   = "FFFFFF"

SERVER = "10.30.21.123"
BASE   = f"http://{SERVER}:8123"
DIRECT = f"http://{SERVER}:9292"

# ── Helpers ─────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def cell_border(cell, hex_color="CCCCCC", sz="4"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),   "single")
        b.set(qn("w:sz"),    sz)
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), hex_color)
        tcBorders.append(b)
    tcPr.append(tcBorders)

def add_heading(text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    if level == 1:
        run.font.size  = Pt(16)
        run.font.bold  = True
        run.font.color.rgb = C_H1
        run.font.name  = "Arial"
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot  = OxmlElement("w:bottom")
        bot.set(qn("w:val"),   "single")
        bot.set(qn("w:sz"),    "6")
        bot.set(qn("w:space"), "4")
        bot.set(qn("w:color"), "2E5496")
        pBdr.append(bot)
        pPr.append(pBdr)
    elif level == 2:
        run.font.size  = Pt(13)
        run.font.bold  = True
        run.font.color.rgb = C_H2
        run.font.name  = "Arial"
    else:
        run.font.size  = Pt(11)
        run.font.bold  = True
        run.font.color.rgb = C_H3
        run.font.name  = "Arial"

def add_para(text="", bold=False, italic=False, size=10.5, color=None, indent=0, space_after=4):
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    run.font.name   = "Arial"
    if color:
        run.font.color.rgb = color
    return p

def add_bullet(text, indent=0.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent   = Cm(indent)
    p.paragraph_format.space_before  = Pt(1)
    p.paragraph_format.space_after   = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = "Arial"

def add_code(text, indent=0):
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  C_BG_CODE)
    pPr  = p._p.get_or_add_pPr()
    pPr.append(shd)

def add_note(text):
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Cm(0.3)
    run = p.add_run(f"  {text}")
    run.font.size   = Pt(10)
    run.font.italic = True
    run.font.name   = "Arial"
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  C_BG_NOTE)
    pPr  = p._p.get_or_add_pPr()
    pPr.append(shd)

def add_method_line(method, path, color, bg):
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    r1 = p.add_run(f"  {method}  ")
    r1.font.name  = "Courier New"
    r1.font.size  = Pt(10)
    r1.font.bold  = True
    r1.font.color.rgb = RGBColor.from_string(C_WHITE)
    # method badge background via highlight is not reliable — use inline run shading via XML
    rPr1 = r1._r.get_or_add_rPr()
    shd1 = OxmlElement("w:shd")
    shd1.set(qn("w:val"),   "clear")
    shd1.set(qn("w:color"), "auto")
    shd1.set(qn("w:fill"),  bg)
    rPr1.append(shd1)

    r2 = p.add_run(f"  {path}")
    r2.font.name  = "Courier New"
    r2.font.size  = Pt(10.5)
    r2.font.bold  = True
    r2.font.color.rgb = color

def add_table(headers, rows, col_widths_cm):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    # Header row
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.width = Cm(col_widths_cm[i])
        set_cell_bg(cell, C_BG_HDR)
        cell_border(cell, "AAAAAA")
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.name = "Arial"
    # Data rows
    for ri, row in enumerate(rows):
        tr = t.rows[ri + 1]
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            cell.width = Cm(col_widths_cm[ci])
            set_cell_bg(cell, "FAFAFA" if ri % 2 == 0 else C_WHITE)
            cell_border(cell, "CCCCCC")
            p = cell.paragraphs[0]
            if isinstance(val, tuple):
                text, bold = val
            else:
                text, bold = val, False
            run = p.add_run(text)
            run.font.bold = bold
            run.font.size = Pt(10)
            run.font.name = "Courier New" if ci == 1 else "Arial"
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
p.paragraph_format.space_after  = Pt(6)
r = p.add_run("HƯỚNG DẪN TEST API")
r.font.size  = Pt(22)
r.font.bold  = True
r.font.color.rgb = C_H1
r.font.name  = "Arial"

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
r = p.add_run("Subscription Management – billing-CA")
r.font.size  = Pt(14)
r.font.color.rgb = C_H2
r.font.name  = "Arial"

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(60)
r = p.add_run("Ngày tạo: 20/05/2026  |  Phiên bản: 1.0")
r.font.size   = Pt(10)
r.font.italic = True
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
r.font.name   = "Arial"

# Info box
t = doc.add_table(rows=3, cols=2)
t.style = "Table Grid"
info_rows = [
    ("Server IP",      SERVER),
    ("Base URL (qua nginx)", f"http://{SERVER}:8123"),
    ("Swagger UI",     f"http://{SERVER}:8123/swagger-ui.html"),
]
for i, (k, v) in enumerate(info_rows):
    c0, c1 = t.rows[i].cells[0], t.rows[i].cells[1]
    c0.width = Cm(5); c1.width = Cm(12)
    set_cell_bg(c0, C_BG_HDR); set_cell_bg(c1, "F8FBFF")
    cell_border(c0); cell_border(c1)
    r0 = c0.paragraphs[0].add_run(k);  r0.font.bold=True; r0.font.size=Pt(10); r0.font.name="Arial"
    r1 = c1.paragraphs[0].add_run(v);  r1.font.size=Pt(10); r1.font.name="Courier New"

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 1. TỔNG QUAN
# ════════════════════════════════════════════════════════════════════════════
add_heading("1. Tổng Quan", 1)
add_para("Tài liệu này hướng dẫn test 5 nhóm API liên quan đến luồng subscription cho khách hàng cá nhân trong hệ thống billing-CA.")
add_para("")

add_table(
    ["Nhóm API", "Base path", "Auth", "Mục đích"],
    [
        ("Public – Gói cước",    "/api/v1/public/plans",          "Không cần",    "Xem danh sách gói hiện hành"),
        ("Public – Subscribe",   "/api/v1/public/subscriptions",  "Không cần",    "Đăng ký + xem subscription"),
        ("Admin – Subscriptions","/api/v1/runtime-subscriptions", "Bearer Token", "Quản lý toàn bộ subscription"),
        ("Admin – Đổi status",   "/api/v1/runtime-subscriptions/{id}/status", "Bearer Token", "Cập nhật trạng thái"),
        ("Admin – Login",        "/api/v1/auth/login",            "Không cần",    "Lấy access token"),
    ],
    [5, 6.5, 2.5, 4.5]
)

# ════════════════════════════════════════════════════════════════════════════
# 2. CÁCH SỬ DỤNG SWAGGER UI
# ════════════════════════════════════════════════════════════════════════════
add_heading("2. Sử Dụng Swagger UI", 1)
add_para("Swagger UI là cách nhanh nhất để test API. Truy cập:")
add_code(f"http://{SERVER}:8123/swagger-ui.html")
add_para("")
add_para("Các bước thao tác trên Swagger:", bold=True)
add_bullet("Mở URL trên trình duyệt Chrome/Firefox")
add_bullet("Với API public: tìm endpoint → bấm Try it out → điền tham số → Execute")
add_bullet("Với API admin: cần login trước (xem mục 3), copy token → bấm nút Authorize (khóa) ở góc trên → dán token vào")
add_note("Lưu ý: Swagger yêu cầu nhập token theo định dạng:  Bearer <access_token>")

# ════════════════════════════════════════════════════════════════════════════
# 3. ĐĂNG NHẬP LẤY TOKEN
# ════════════════════════════════════════════════════════════════════════════
add_heading("3. Đăng Nhập Lấy Access Token (Admin APIs)", 1)
add_para("Các API admin yêu cầu JWT Bearer Token. Thực hiện đăng nhập trước:")

add_heading("3.1  Endpoint", 2)
add_method_line("POST", f"http://{SERVER}:8123/api/v1/auth/login", C_BLUE, "0B4F8A")

add_heading("3.2  Request Body", 2)
add_code('Content-Type: application/json')
add_code('')
add_code('{')
add_code('  "username": "admin",')
add_code('  "password": "Admin@123"')
add_code('}')

add_heading("3.3  Response thành công (200)", 2)
add_code('{')
add_code('  "data": {')
add_code('    "accessToken": "eyJhbGciOiJSUzI1NiJ9...",   // dùng token này')
add_code('    "refreshToken": "...",')
add_code('    "expiresIn": 3600')
add_code('  }')
add_code('}')

add_note("Sao chép giá trị accessToken. Token có hiệu lực 1 giờ.")

add_heading("3.4  Gắn Token Vào Request", 2)
add_para("Khi dùng Postman/curl, thêm header:")
add_code('Authorization: Bearer eyJhbGciOiJSUzI1NiJ9...')
add_para("Khi dùng Swagger UI: bấm Authorize → nhập  Bearer <token>")

# ════════════════════════════════════════════════════════════════════════════
# 4. PUBLIC API – DANH SÁCH GÓI CƯỚC
# ════════════════════════════════════════════════════════════════════════════
add_heading("4. API – Lấy Danh Sách Gói Cước Hiện Hành", 1)
add_note("Không cần đăng nhập. Gọi trực tiếp từ trình duyệt hoặc Postman.")

add_heading("4.1  Endpoint", 2)
add_method_line("GET", f"http://{SERVER}:8123/api/v1/public/plans", C_GREEN, "1A6B2E")

add_heading("4.2  Không có tham số", 2)

add_heading("4.3  Response thành công (200)", 2)
add_code('[')
add_code('  {')
add_code('    "scheduleId": 1,              // dùng để subscribe ở API tiếp theo')
add_code('    "subjectType": "INDIVIDUAL",  // CÁ NHÂN')
add_code('    "planName": "Gói chữ ký số cá nhân",')
add_code('    "iconUrl": "http://...",')
add_code('    "minFee": 30000,')
add_code('    "minFeeFormatted": "30.000đ",')
add_code('    "features": ["Ký hợp đồng điện tử", "Khai báo thuế", ...]')
add_code('  },')
add_code('  { "subjectType": "ORGANIZATION", ... },')
add_code('  { "subjectType": "INDIVIDUAL_OF_ORG", ... }')
add_code(']')

add_heading("4.4  Lưu Ý", 2)
add_bullet("scheduleId trả về là ID của đợt gói đang ACTIVE – cần lưu lại để dùng cho API đăng ký bên dưới")
add_bullet("Nếu response là [] (mảng rỗng) → chưa có RetailPlanSchedule nào ở trạng thái ACTIVE hoặc APPROVED")

# ════════════════════════════════════════════════════════════════════════════
# 5. PUBLIC API – ĐĂNG KÝ GÓI
# ════════════════════════════════════════════════════════════════════════════
add_heading("5. API – Đăng Ký Gói (Subscribe)", 1)
add_note("Không cần đăng nhập. Gọi sau khi người dùng thanh toán thành công.")

add_heading("5.1  Endpoint", 2)
add_method_line("POST", f"http://{SERVER}:8123/api/v1/public/subscriptions", C_BLUE, "0B4F8A")

add_heading("5.2  Request Body", 2)
add_table(
    ["Trường", "Kiểu", "Bắt buộc", "Mô tả"],
    [
        ("userId",           "Long",    "Có",     "ID khách hàng từ hệ thống RS Core"),
        ("scheduleId",       "Long",    "Có",     "Lấy từ kết quả API /public/plans"),
        ("pricingRuleId",    "Long",    "Không",  "Cụ thể bậc giá muốn chọn, bỏ trống = tự chọn bậc thấp nhất"),
        ("paymentReference", "String",  "Không",  "Mã giao dịch thanh toán (VNPay, MoMo...)"),
    ],
    [3.5, 2, 2, 9]
)
add_code('{')
add_code('  "userId": 1001,')
add_code('  "scheduleId": 1,')
add_code('  "pricingRuleId": 2,          // tùy chọn')
add_code('  "paymentReference": "VNP_TXN_20260520_ABC123"  // tùy chọn')
add_code('}')

add_heading("5.3  Response thành công (200)", 2)
add_code('{')
add_code('  "code": "0000",')
add_code('  "message": "Subscription created successfully",')
add_code('  "data": {')
add_code('    "subscriptionId": 42,       // lưu lại ID này')
add_code('    "subscriberType": "INDIVIDUAL",')
add_code('    "userId": 1001,')
add_code('    "planTemplateId": 1,')
add_code('    "planCode": "PKG-CA-INDIVIDUAL",')
add_code('    "planName": "Gói chữ ký số cá nhân",')
add_code('    "status": "ACTIVE",')
add_code('    "startDate": "2026-05-20",')
add_code('    "endDate": "2027-05-20",')
add_code('    "signingQuotaTotal": 50,')
add_code('    "signingQuotaUsed": 0,')
add_code('    "paymentReference": "VNP_TXN_20260520_ABC123"')
add_code('  }')
add_code('}')

add_heading("5.4  Lỗi thường gặp", 2)
add_table(
    ["HTTP Code", "Mô tả", "Nguyên nhân"],
    [
        ("409", "Subscription đã tồn tại",     "userId đã subscribe scheduleId này và chưa hết hạn"),
        ("400", "Validation failed",           "Thiếu userId hoặc scheduleId"),
        ("409", "Schedule chưa ACTIVE",        "RetailPlanSchedule chưa được admin kích hoạt"),
        ("404", "Schedule not found",          "scheduleId không tồn tại"),
    ],
    [2.5, 4.5, 9.5]
)

# ════════════════════════════════════════════════════════════════════════════
# 6. PUBLIC API – XEM SUBSCRIPTION THEO USER
# ════════════════════════════════════════════════════════════════════════════
add_heading("6. API – Xem Danh Sách Subscription Theo userId", 1)
add_note("Không cần đăng nhập.")

add_heading("6.1  Endpoint", 2)
add_method_line("GET", f"http://{SERVER}:8123/api/v1/public/subscriptions?userId=1001", C_GREEN, "1A6B2E")

add_heading("6.2  Query Parameter", 2)
add_table(
    ["Tham số", "Kiểu", "Bắt buộc", "Mô tả"],
    [("userId", "Long", "Có", "ID khách hàng từ hệ thống RS Core")],
    [3, 2, 2, 9.5]
)

add_heading("6.3  Response thành công (200)", 2)
add_code('{')
add_code('  "code": "0000",')
add_code('  "data": [')
add_code('    {')
add_code('      "subscriptionId": 42,')
add_code('      "status": "ACTIVE",')
add_code('      "planName": "Gói chữ ký số cá nhân",')
add_code('      "startDate": "2026-05-20",')
add_code('      "endDate": "2027-05-20",')
add_code('      "signingQuotaTotal": 50,')
add_code('      "signingQuotaUsed": 3')
add_code('    }')
add_code('  ]')
add_code('}')

# ════════════════════════════════════════════════════════════════════════════
# 7. ADMIN API – QUẢN LÝ SUBSCRIPTION
# ════════════════════════════════════════════════════════════════════════════
add_heading("7. Admin API – Quản Lý Subscription (Cần Auth)", 1)
add_note("Yêu cầu Bearer Token. Thực hiện bước đăng nhập ở mục 3 trước.")

add_heading("7.1  Lấy Tất Cả Subscription", 2)
add_method_line("GET", f"http://{SERVER}:8123/api/v1/runtime-subscriptions", C_GREEN, "1A6B2E")
add_para("Trả về toàn bộ subscription, sắp xếp theo ngày tạo mới nhất.", indent=0.3)

add_heading("7.2  Lấy Subscription Theo ID", 2)
add_method_line("GET", f"http://{SERVER}:8123/api/v1/runtime-subscriptions/{{id}}", C_GREEN, "1A6B2E")
add_para("Thay {id} bằng subscriptionId cụ thể, ví dụ: .../runtime-subscriptions/42", indent=0.3)

add_heading("7.3  Cập Nhật Trạng Thái Subscription", 2)
add_method_line("PATCH", f"http://{SERVER}:8123/api/v1/runtime-subscriptions/{{id}}/status", C_ORANGE, "B45A00")

add_para("Request Body:", bold=True, indent=0.3)
add_code('{')
add_code('  "status": "SUSPENDED",    // trạng thái muốn chuyển sang')
add_code('  "actor": "nguyen.van.a"   // tên người thực hiện (ghi vào audit log)')
add_code('}')

add_para("Các chuyển trạng thái hợp lệ:", bold=True, indent=0.3)
add_table(
    ["Từ (current)", "Sang được (next)", "Ghi chú"],
    [
        ("PENDING",   "ACTIVE, CANCELLED",                  "Subscription mới tạo"),
        ("ACTIVE",    "SUSPENDED, CANCELLED, EXPIRED",      "Đang hoạt động"),
        ("SUSPENDED", "ACTIVE, CANCELLED, EXPIRED",         "Tạm dừng"),
        ("EXPIRED",   "(không chuyển được)",                "Trạng thái kết thúc"),
        ("CANCELLED", "(không chuyển được)",                "Trạng thái kết thúc"),
    ],
    [3.5, 5.5, 7.5]
)
add_note("Nếu chuyển sai (ví dụ EXPIRED → ACTIVE) sẽ nhận HTTP 409 với message lỗi rõ ràng.")

# ════════════════════════════════════════════════════════════════════════════
# 8. KỊCH BẢN TEST ĐẦU ĐỦ
# ════════════════════════════════════════════════════════════════════════════
add_heading("8. Kịch Bản Test Đề Xuất", 1)

add_heading("Kịch bản 1 – Luồng đầy đủ khách hàng cá nhân", 2)
steps = [
    ("Bước 1", "GET /api/v1/public/plans",
     "Lấy danh sách gói, copy scheduleId từ kết quả"),
    ("Bước 2", "POST /api/v1/public/subscriptions",
     "Đăng ký với userId=1001, scheduleId từ bước 1 → nhận subscriptionId"),
    ("Bước 3", "GET /api/v1/public/subscriptions?userId=1001",
     "Xác nhận subscription đã tạo, status=ACTIVE"),
    ("Bước 4", "POST /api/v1/auth/login",
     "Đăng nhập admin, lấy accessToken"),
    ("Bước 5", "GET /api/v1/runtime-subscriptions/{id}",
     "Admin xem chi tiết subscription vừa tạo"),
    ("Bước 6", "PATCH /api/v1/runtime-subscriptions/{id}/status",
     'Admin đổi status sang SUSPENDED với body {"status":"SUSPENDED","actor":"admin"}'),
    ("Bước 7", "PATCH /api/v1/runtime-subscriptions/{id}/status",
     'Thử đổi EXPIRED → ACTIVE → phải nhận HTTP 409'),
]
add_table(
    ["Bước", "API", "Mục tiêu"],
    steps,
    [2, 6, 8.5]
)

add_heading("Kịch bản 2 – Kiểm tra chặn duplicate subscription", 2)
add_bullet("Gọi POST /api/v1/public/subscriptions lần 2 với cùng userId + scheduleId")
add_bullet("Kỳ vọng: trả về subscription cũ (không tạo mới), HTTP 200")

add_heading("Kịch bản 3 – Kiểm tra transition không hợp lệ", 2)
add_bullet("Tạo subscription (status=ACTIVE)")
add_bullet("Gọi PATCH đổi sang EXPIRED → HTTP 200 (hợp lệ)")
add_bullet("Gọi PATCH đổi sang ACTIVE → HTTP 409 (không hợp lệ từ EXPIRED)")

# ════════════════════════════════════════════════════════════════════════════
# 9. CURL REFERENCE
# ════════════════════════════════════════════════════════════════════════════
add_heading("9. Tham Khảo – Lệnh curl", 1)

add_heading("Lấy danh sách gói:", 2)
add_code(f'curl -X GET "http://{SERVER}:8123/api/v1/public/plans"')

add_heading("Đăng ký gói:", 2)
add_code(f'curl -X POST "http://{SERVER}:8123/api/v1/public/subscriptions" \\')
add_code('  -H "Content-Type: application/json" \\')
add_code('  -d \'{"userId":1001,"scheduleId":1,"paymentReference":"TEST-001"}\'')

add_heading("Lấy subscription theo userId:", 2)
add_code(f'curl -X GET "http://{SERVER}:8123/api/v1/public/subscriptions?userId=1001"')

add_heading("Đăng nhập admin:", 2)
add_code(f'curl -X POST "http://{SERVER}:8123/api/v1/auth/login" \\')
add_code('  -H "Content-Type: application/json" \\')
add_code('  -d \'{"username":"admin","password":"Admin@123"}\'')

add_heading("Đổi trạng thái subscription:", 2)
add_code(f'curl -X PATCH "http://{SERVER}:8123/api/v1/runtime-subscriptions/42/status" \\')
add_code('  -H "Content-Type: application/json" \\')
add_code('  -H "Authorization: Bearer <access_token>" \\')
add_code('  -d \'{"status":"SUSPENDED","actor":"nguyen.van.a"}\'')

# ════════════════════════════════════════════════════════════════════════════
# SAVE
# ════════════════════════════════════════════════════════════════════════════
out = r"E:\PhuongCM\billing-CA\docs\Huong_dan_test_API_subscription.docx"
doc.save(out)
print(f"Saved: {out}")
